from docx import Document
import tempfile
import re

def fill_placeholders(file_bytes: bytes, responses):
    print("\n==============================")
    print("🧾 Starting fill_placeholders()")
    print("Responses received:", responses)
    print("==============================\n")

    # Handle both new format (ordered list) and legacy format (dict)
    if isinstance(responses, list):
        # New format: ordered list of {id, label, value}
        ordered_values = [item.get("value", "") for item in responses]
        is_ordered_format = True
        print(f"✅ Using ordered format with {len(ordered_values)} values")
    else:
        # Legacy format: dict mapping labels to values
        label_map = responses
        is_ordered_format = False
        print("✅ Using legacy dict format")
    
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        tmp.write(file_bytes)
        tmp_path = tmp.name
        print(f"📁 Temporary file created at: {tmp_path}")

    doc = Document(tmp_path)
    print("✅ Document loaded successfully.")
    
    # Track which occurrence we're on (for ordered format)
    occurrence_index = [0]
    
    # Placeholder patterns (same as parser uses)
    PLACEHOLDER_PATTERNS = [
        (r"\[[^\]]+\]", r"\[([^\]]+)\]"),           # [KEY]
        (r"\{\{[^}]+\}\}", r"\{\{([^}]+)\}\}"),     # {{KEY}}
        (r"<[^>]+>", r"<([^>]+)>"),                  # <KEY>
        (r"\$\s*\[[^\]]+\]", r"\$\s*\[([^\]]+)\]"), # $[KEY]
    ]
    
    def replace_in_paragraph(paragraph):
        """Replace placeholders in paragraph, handling both formats."""
        full_text = "".join(run.text for run in paragraph.runs)
        replaced_text = full_text
        
        if is_ordered_format:
            # Ordered format: find placeholders and replace in order
            # Find all placeholders in this paragraph
            matches = []
            for pattern, extract_pattern in PLACEHOLDER_PATTERNS:
                for match in re.finditer(pattern, full_text):
                    matches.append((match.start(), match.end(), match.group(0), pattern))
            
            # Sort by position (left to right)
            matches.sort(key=lambda x: x[0])
            
            # Replace in reverse order to preserve positions
            for start, end, raw, _ in reversed(matches):
                if occurrence_index[0] < len(ordered_values):
                    value = ordered_values[occurrence_index[0]]
                    if value:
                        replaced_text = replaced_text[:start] + str(value) + replaced_text[end:]
                        print(f"🔁 Replacing occurrence {occurrence_index[0]} '{raw}' with '{value}'")
                    occurrence_index[0] += 1
        else:
            # Legacy format: label-based replacement
            # Normal bracketed placeholders
            for key, value in label_map.items():
                if key not in ("$[__________]", "_____________") and value:
                    patterns = [
                        rf"\[\s*{re.escape(key)}\s*\]",   # [KEY]
                        rf"\{{\s*{re.escape(key)}\s*\}}", # {{KEY}}
                    ]
                    for pattern in patterns:
                        if re.search(pattern, replaced_text, flags=re.IGNORECASE):
                            replaced_text = re.sub(pattern, str(value), replaced_text, flags=re.IGNORECASE)
                            print(f"🔁 Replacing placeholder '{key}' with '{value}'")
            
            # Money placeholders
            money_value = label_map.get("$[__________]")
            if money_value:
                if re.search(r"\$\s*\[[^\]]+\]", replaced_text):
                    replaced_text = re.sub(
                        r"\$\s*\[[^\]]+\]",
                        f"${money_value}",
                        replaced_text
                    )
                    print(f"💰 Replacing money placeholder with '${money_value}'")
            
            # Underline blanks
            underline_value = label_map.get("_____________")
            if underline_value:
                if re.search(r"_{3,}", replaced_text):
                    replaced_text = re.sub(r"_{3,}", str(underline_value), replaced_text)
                    print(f"🖊️ Replacing underline blanks with '{underline_value}'")
        
        # Write back to paragraph
        if full_text != replaced_text:
            for run in paragraph.runs:
                run.text = ""
            if paragraph.runs:
                paragraph.runs[0].text = replaced_text
            else:
                paragraph.add_run(replaced_text)
    
    def process_element(element):
        """Recursively process all paragraphs and tables."""
        if hasattr(element, "paragraphs"):
            for p in element.paragraphs:
                replace_in_paragraph(p)
        if hasattr(element, "tables"):
            for t in element.tables:
                for r in t.rows:
                    for c in r.cells:
                        process_element(c)
    
    # Process all paragraphs and tables in the main body
    process_element(doc)
    
    # Process headers and footers
    for section in doc.sections:
        if section.header:
            process_element(section.header)
        if section.footer:
            process_element(section.footer)

    # Save filled file
    output_path = tmp_path.replace(".docx", "_filled.docx")
    doc.save(output_path)
    print(f"✅ Document saved successfully at: {output_path}")
    print(f"📊 Processed {occurrence_index[0]} occurrences" if is_ordered_format else "")
    print("==============================\n")

    return output_path
